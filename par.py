from bale import Bot, CallbackQuery, Message, InlineKeyboardMarkup, InlineKeyboardButton, MenuKeyboardMarkup, MenuKeyboardButton
import chess
import chess.engine
import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from reportlab.pdfgen import canvas
from pathlib import Path
import arabic_reshaper
from bidi.algorithm import get_display
import re
import translate1111

# 📂 مسیر به موتور Stockfish
STOCKFISH_PATH = r"par/stockfish-windows-x86-64-sse41-popcnt.exe"

# ♟️ نمادهای یونیکد برای مهره‌های شطرنج
PIECE_SYMBOLS = {
    'p': '♙', 'r': '♖', 'n': '♘', 'b': '♗', 'q': '♕', 'k': '♔',
    'P': '♟', 'R': '♜', 'N': '♞', 'B': '♝', 'Q': '♛', 'K': '♚'
}

# 🤖 راه‌اندازی ربات با توکن شما
client = Bot(token="1073138097:49XMJIzGSAXXffU4p5hFbiDQB6NLp7RCxqdDpPeW")
user_input_state = {}

def pdf_to_docx(pdf_path, docx_path):
    """🔄 تبدیل PDF به DOCX"""
    reader = PdfReader(pdf_path)
    document = Document()

    for page in reader.pages:
        text = page.extract_text()
        text = reshape_text_if_persian(text)
        document.add_paragraph(text)

    document.save(docx_path)
    print(f"📄 تبدیل {pdf_path} به {docx_path}")

def docx_to_pdf(docx_path, pdf_path):
    """🔄 تبدیل DOCX به PDF"""
    document = Document(docx_path)
    pdf = canvas.Canvas(pdf_path)

    y = 800 
    for i, paragraph in enumerate(document.paragraphs):
        text = reshape_text_if_persian(paragraph.text)
        pdf.drawString(100, y - 15 * i, text)

    pdf.save()
    print(f"📄 تبدیل {docx_path} به {pdf_path}")

def is_persian(text):
    """📝 بررسی اینکه آیا متن شامل کاراکترهای فارسی است"""
    return bool(re.search(r'[\u0600-\u06FF]', text))

def reshape_text_if_persian(text):
    """🔤 بازشکل‌دهی متن اگر به فارسی است"""
    if is_persian(text):
        reshaped_text = arabic_reshaper.reshape(text)
        return get_display(reshaped_text)
    return text

def display_board(board):
    """📊 نمایش صفحه شطرنج"""
    board = [row.split() for row in board]
    fen = ''
    for row in board:
        empty = 0
        for cell in row:
            if cell == '.':
                empty += 1
            else:
                if empty > 0:
                    fen += str(empty)
                    empty = 0
                fen += cell
        if empty > 0:
            fen += str(empty)
        fen += '/'
    fen = fen[:-1]
    fen = f"https://lichess1.org/export/fen.gif?fen={fen}&color=white"
    return fen


async def play_chess(chat_id, color):
    """♟️ بازی شطرنج با کاربر"""
    if not os.path.exists(STOCKFISH_PATH):
        await client.send_message(chat_id=chat_id, text="❌ موتور Stockfish یافت نشد! لطفاً مسیر را بررسی کنید.")
        return

    board = chess.Board()
    user_input_state[chat_id] = {"board": board, "color": color}

    boardFen = f"https://lichess1.org/export/fen.gif?fen={board.fen().replace(" ", "_")}&color=white"

    try:
        with chess.engine.SimpleEngine.popen_uci(STOCKFISH_PATH) as engine:
            if color == "black":
                result = engine.play(board, chess.engine.Limit(time=1.0))
                board.push(result.move)
                await client.send_message(chat_id=chat_id, text=f"♟️ **Stockfish بازی را شروع کرد!**\n{boardFen}\n♔ **نوبت شما! حرکت خود را ارسال کنید.**")
            else:
                await client.send_message(chat_id=chat_id, text=f"♔ **شما سفید هستید! بازی را شروع کنید.**\n{boardFen}")

            user_input_state[chat_id]["awaiting_move"] = True
    except PermissionError as e:
        await client.send_message(chat_id=chat_id, text="❌ خطای دسترسی: لطفاً مجوزهای اجرای Stockfish را بررسی کنید.")
    except Exception as e:
        await client.send_message(chat_id=chat_id, text=f"❌ خطای غیرمنتظره: {e}")

@client.event
async def on_message(message: Message):
    chat_id = message.chat.id
    if chat_id in user_input_state and user_input_state[chat_id].get("awaiting"):
        if user_input_state[chat_id]["awaiting"] == "origin":
            user_input_state[chat_id]["origin"] = message.content
            user_input_state[chat_id]["awaiting"] = "dest"
            await client.send_message(chat_id=chat_id, text="🌐 لطفاً زبان مقصد را وارد کنید:")
        elif user_input_state[chat_id]["awaiting"] == "dest":
            user_input_state[chat_id]["dest"] = message.content
            await client.send_message(chat_id=chat_id, text="📝 لطفاً متن برای ترجمه را وارد کنید:")
            user_input_state[chat_id]["awaiting"] = "text"
        elif user_input_state[chat_id]["awaiting"] == "text":
            user_input_state[chat_id]["text"] = message.content
            origin = user_input_state[chat_id]["origin"]
            dest = user_input_state[chat_id]["dest"]
            text = user_input_state[chat_id]["text"]
            
            try:
                translation = translate1111.translate(text, origin, dest)
                await client.send_message(chat_id=chat_id, text=f"🔠 ترجمه: {translation}")
            except Exception as e:
                await client.send_message(chat_id=chat_id, text=f"❌ خطا در ترجمه: {e}")
            
            user_input_state[chat_id]["awaiting"] = None
        return
    
    if message.content == "/start":
        reply_markup = InlineKeyboardMarkup()
        reply_markup.add(InlineKeyboardButton(text="📄 تبدیل PDF به DOCX", callback_data="pdf_to_docx"), row=1)
        reply_markup.add(InlineKeyboardButton(text="🔠 ترجمه متن", callback_data="translate_text"), row=2)
        reply_markup.add(InlineKeyboardButton(text="♟️ بازی شطرنج", callback_data="chess"), row=3)
        reply_markup.add(InlineKeyboardButton(text="🎥 دانلود ویدئو یوتیوب", callback_data="download_youtube"), row=4)
        reply_markup.add(InlineKeyboardButton(text="🎤 تبدیل صدا به متن", callback_data="audio_to_text"), row=5)
        await client.send_message(chat_id=chat_id, text=f"*سلام {message.author.first_name}, به IntelliBot خوش آمدید*", components=reply_markup)

    elif message.content == "/keyboard":
        await client.send_message(chat_id=chat_id, text=f"*سلام {message.author.first_name}, به IntelliBot خوش آمدید*",
                            components=MenuKeyboardMarkup().add(MenuKeyboardButton('🔠 ترجمه')).add(MenuKeyboardButton('📄 pdf docx')))

    elif message.content == "ترجمه":
        user_input_state[chat_id] = {"awaiting": "origin"}
        await client.send_message(chat_id=chat_id, text="🌐 لطفاً زبان مبدأ را وارد کنید:")

    else:
        await client.send_message(chat_id=chat_id, text=f"شما گفتید: {translate1111.translate}")

    print("📨 پیام دریافت شد")
    if message.document:
        file = message.document
        file_path = f'./{file.file_name}'

        await file.download(file_path)
        
        if file.file_name.endswith('.pdf'):
            docx_path = f"{file_path.replace('.pdf', '.docx')}"
            pdf_to_docx(file_path, docx_path)
            await client.send_document(chat_id=chat_id, document=docx_path)
            os.remove(file_path)
            os.remove(docx_path)
        
        elif file.file_name.endswith('.docx'):
            pdf_path = f"{file_path.replace('.docx', '.pdf')}"
            docx_to_pdf(file_path, pdf_path)
            await client.send_document(chat_id=chat_id, document=pdf_path)
            os.remove(file_path)
            os.remove(pdf_path)

        else:
            await client.send_message(chat_id=chat_id, text="❌ فقط فایل‌های PDF یا DOCX قابل تبدیل هستند.")
    
    elif chat_id in user_input_state and user_input_state[chat_id].get("awaiting_move"):
        board = user_input_state[chat_id]["board"]
        try:
            move = board.parse_san(message.content)
            if move in board.legal_moves:
                board.push(move)
            else:
                await client.send_message(chat_id=chat_id, text="❌ حرکت نامعتبر. دوباره امتحان کنید.")
                return
        except ValueError:
            await client.send        .send_message(chat_id=chat_id, text="❌ فرمت حرکت نامعتبر. دوباره امتحان کنید.")
            return

        with chess.engine.SimpleEngine.popen_uci(STOCKFISH_PATH) as engine:
            if board.is_game_over():
                await client.send_message(chat_id=chat_id, text=f"🎉 بازی تمام شد! نتیجه: {board.result()}")
                del user_input_state[chat_id]
                return

            result = engine.play(board, chess.engine.Limit(time=1.0))
            board.push(result.move)

            if board.is_game_over():
                await client.send_message(chat_id=chat_id, text=f"🎉 بازی تمام شد! نتیجه: {board.result()}")
                del user_input_state[chat_id]
            else:
                await client.send_message(chat_id=chat_id, text=f"♟️ **Stockfish حرکت کرد!**\n{display_board(board)}\n♔ **نوبت شماست! حرکت خود را بفرستید.**")

@client.event
async def on_ready():
    print(client.user, "آماده است! 🟢")

@client.event
async def on_callback(callback: CallbackQuery):
    chat_id = callback.message.chat.id
    if callback.data == "chess":
        keyboard = InlineKeyboardMarkup()
        keyboard.add(InlineKeyboardButton(text="⚪ سفید", callback_data="chess_white"))
        keyboard.add(InlineKeyboardButton(text="⚫ سیاه", callback_data="chess_black"))
        await client.send_message(chat_id=chat_id, text="🎲 **با چه رنگی می‌خواهید بازی کنید؟**", components=keyboard)
    elif callback.data == "chess_white":
        await play_chess(chat_id, "white")
    elif callback.data == "chess_black":
        await play_chess(chat_id, "black")
    elif callback.data == "translate_text":
        user_input_state[chat_id] = {"awaiting": "origin"}
        await client.send_message(chat_id=chat_id, text="🌐 لطفاً زبان مبدأ را وارد کنید:")

client.run()
