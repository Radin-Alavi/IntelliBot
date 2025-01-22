import re
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from pathlib import Path
import pdfplumber

def extract_text_from_pdf(pdf_path):
    """استخراج متن از PDF با استفاده از pdfplumber با حفظ خط‌های اصلی"""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                # حذف کاراکترهای غیرمجاز
                page_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]+', '', page_text)
                text += page_text + "\n"
    return text

def pdf_to_docx(pdf_path, docx_path):
    """تبدیل PDF به DOCX"""
    document = Document()
    text = extract_text_from_pdf(pdf_path)
    document.add_paragraph(text)
    document.save(docx_path)
    print(f"تبدیل شد {pdf_path} به {docx_path}")

def docx_to_pdf(docx_path, pdf_path):
    """تبدیل DOCX به PDF"""
    document = Document(docx_path)
    pdf = canvas.Canvas(pdf_path)

    for i, paragraph in enumerate(document.paragraphs):
        # تنظیم موقعیت متن برای راست‌به‌چپ
        pdf.drawRightString(500, 800 - 15 * i, paragraph.text)

    pdf.save()
    print(f"تبدیل شد {docx_path} به {pdf_path}")

def main():
    print("یک گزینه را انتخاب کنید")
    print("1. تبدیل PDF به DOCX")
    print("2. تبدیل DOCX به PDF")
    choice = input(" انتخاب کنید (1/2) ").strip()

    if choice == "1":
        pdf_path = input("مسیر فایل PDF را وارد کنید ").strip()
        docx_path = input("مسیر فایل DOCX را وارد کنید  ").strip()
        if Path(pdf_path).exists():
            pdf_to_docx(pdf_path, docx_path)
        else:
            print("این فایل PDF وجود ندارد")

    elif choice == "2":
        docx_path = input("مسیر فایل DOCX را وارد کنید  ").strip()
        pdf_path = input("مسیر فایل PDF را وارد کنید  ").strip()
        if Path(docx_path).exists():
            docx_to_pdf(docx_path, pdf_path)
        else:
            print("این فایل DOCX وجود ندارد")

    else:
        print("از بین 1 و 2 انتخاب کنید")

if __name__ == "__main__":
    main()
