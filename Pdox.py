from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from reportlab.pdfgen import canvas
from pathlib import Path
import arabic_reshaper
from bidi.algorithm import get_display
import re

def is_persian(text):
    return bool(re.search(r'[\u0600-\u06FF]', text))

def reshape_text_if_persian(text):
    if is_persian(text):
        reshaped_text = arabic_reshaper.reshape(text)
        return get_display(reshaped_text)
    return text

def pdf_to_docx(pdf_path, docx_path):
    """تبدیل PDF به DOCX"""
    reader = PdfReader(pdf_path)
    document = Document()

    for page in reader.pages:
        text = page.extract_text()
        text = reshape_text_if_persian(text)
        document.add_paragraph(text)

    document.save(docx_path)
    print(f"تبدیل شد {pdf_path} به {docx_path}")

def docx_to_pdf(docx_path, pdf_path):
    """تبدیل DOCX به PDF"""
    document = Document(docx_path)
    pdf = canvas.Canvas(pdf_path)

    y = 800 
    for i, paragraph in enumerate(document.paragraphs):
        text = reshape_text_if_persian(paragraph.text)
        pdf.drawString(100, y - 15 * i, text)

    pdf.save()
    print(f"تبدیل شد {docx_path} به {pdf_path}")

def main():
    print("یک گزینه را انتخاب کنید")
    print("1. تبدیل PDF به DOCX")
    print("2. تبدیل DOCX به PDF")
    choice = input("انتخاب کنید (1/2): ").strip()

    if choice == "1":
        pdf_path = input("مسیر فایل PDF را وارد کنید: ").strip()
        docx_path = input("مسیر فایل DOCX را وارد کنید: ").strip()
        if Path(pdf_path).exists():
            pdf_to_docx(pdf_path, docx_path)
        else:
            print("این فایل PDF وجود ندارد")

    elif choice == "2":
        docx_path = input("مسیر فایل DOCX را وارد کنید: ").strip()
        pdf_path = input("مسیر فایل PDF را وارد کنید: ").strip()
        if Path(docx_path).exists():
            docx_to_pdf(docx_path, pdf_path)
        else:
            print("این فایل DOCX وجود ندارد")

    else:
        print("از بین 1 و 2 انتخاب کنید")

if __name__ == "__main__":
    main()
